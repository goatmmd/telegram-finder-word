import sys
import time

import docx


class DocxWriter:
    def __init__(self, data, target_name, word):
        self.raw_data = data
        self.target_name = target_name
        self.target_word = word
        self.docx_writer: docx.Document = docx.Document()

    def write_to_file(self):
        self.docx_writer.add_heading(f'FINDING WORD {self.target_word} FROM {self.target_name}\n', 1)

        for data in self.raw_data:
            self.docx_writer.add_paragraph(f'**Message**: \n[ {data["text"]} ]\n**Date**: \n{data["date"]}\n')

        if sys.platform == 'linux':
            self.docx_writer.save(f'output/{int(time.time())}.docx')
        else:
            self.docx_writer.save(f'output\\{int(time.time())}.docx')

        print('Your date have been received in output')
